import requests
import json
from docx import Document

def generate_content(endpoint, prompt):
    url = f'http://127.0.0.1:11434/api/{endpoint}'
    headers = {
        'Content-Type': 'application/json'
    }
    data = {
        'model': 'llama3',
        'prompt': prompt
    }
    response = requests.post(url, json=data, headers=headers, stream=True)
    
    if response.status_code == 200:
        content = ''
        for line in response.iter_lines():
            if line:
                json_line = line.decode('utf-8')
                try:
                    result = json.loads(json_line)
                    if result['done']:
                        break
                    content += result.get('response', '')
                except json.JSONDecodeError:
                    print(f"Error decoding JSON: {json_line}")
        return content
    else:
        raise Exception(f"Error: {response.status_code} - {response.text}")

def save_to_docx(content, filename):
    doc = Document()
    doc.add_heading('Generación de Contenido con IA', level=1)
    doc.add_paragraph(content)
    doc.save(filename)
    print(f"Archivo guardado como {filename}")

def generate_html_blog(content, topic, keywords):
    try:
        # Genera el HTML usando el mismo endpoint con un prompt específico
        html_prompt = f"Genera una plantilla HTML para un blog sobre el tema '{topic}'. Incluye las siguientes palabras clave: {', '.join(keywords)}. Usa las mejores prácticas de SEO."
        template = generate_content('generate', html_prompt)
        
        # Reemplaza los placeholders en la plantilla HTML con el contenido real
        html_content = template.replace('{{ topic }}', topic)
        html_content = html_content.replace('{{ keywords }}', ', '.join(keywords))
        html_content = html_content.replace('{{ content }}', content)
        
        with open('Blog_Content.html', 'w', encoding='utf-8') as file:
            file.write(html_content)
        print("Archivo HTML guardado como Blog_Content.html")
    except Exception as e:
        print(f"Error al generar el HTML: {e}")

def generate_social_media_posts(topic):
    try:
        # Genera entradas para redes sociales usando el mismo endpoint con un prompt específico
        social_media_prompt = f"Genera 5 entradas para redes sociales sobre el tema '{topic}'. Incluye hashtags y llamados a la acción. en español"
        posts = generate_content('generate', social_media_prompt)
        
        # Guarda las entradas en un archivo de texto
        with open('Social_Media_Posts.txt', 'w', encoding='utf-8') as file:
            file.write(posts)
        print("Archivo de entradas para redes sociales guardado como Social_Media_Posts.txt")
    except Exception as e:
        print(f"Error al generar las entradas para redes sociales: {e}")

def generate_blog_content(topic, keywords, max_paragraphs):
    url = 'http://127.0.0.1:11434/api/generate'
    headers = {
        'Content-Type': 'application/json'
    }
    prompt = f"Escribe un blog sobre el tema: {topic}. Incluye las siguientes palabras clave: {', '.join(keywords)}. Limita el contenido a {max_paragraphs} párrafos."
    data = {
        'model': 'llama3',
        'prompt': prompt
    }
    response = requests.post(url, json=data, headers=headers, stream=True)
    
    if response.status_code == 200:
        blog_content = ''
        for line in response.iter_lines():
            if line:
                json_line = line.decode('utf-8')
                try:
                    result = json.loads(json_line)
                    if result['done']:
                        break
                    blog_content += result.get('response', '')
                except json.JSONDecodeError:
                    print(f"Error decoding JSON: {json_line}")
        
        # Limitar el contenido a max_paragraphs párrafos
        paragraphs = blog_content.split('\n\n')  # Asumiendo que los párrafos están separados por dos saltos de línea
        if len(paragraphs) > max_paragraphs:
            blog_content = '\n\n'.join(paragraphs[:max_paragraphs])
        
        return blog_content
    else:
        raise Exception(f"Error: {response.status_code} - {response.text}")

# Parámetros del blog
filename ="la Agricultura de Precisión"
topic = "El Futuro de la Agricultura de Precisión: Tecnología y Sostenibilidad en el Cultivo de Alimentos"
keywords = ["Agricultura de precisión", "Tecnología agrícola", "Sostenibilidad", "Sensores de cultivo", "Drones en agricultura", "Big Data en agricultura", "Riego eficiente", "Control de plagas", "Cultivos inteligentes", "Gestión de recursos"]
max_paragraphs = 20  # Limita el blog a 10 párrafos

# Generar el contenido del blog
try:
    blog_content = generate_blog_content(topic, keywords, max_paragraphs)
    save_to_docx(blog_content, f'{filename}.docx')
    generate_html_blog(blog_content, topic, keywords)
    generate_social_media_posts(topic)
except Exception as e:
    print(e)
